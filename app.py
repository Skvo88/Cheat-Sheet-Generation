import streamlit as st
import re
from io import BytesIO
from docx import Document
from docx.shared import Cm, Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# =====================================================================
# БЛОК 1: XML-ПОМОЩНИКИ (СЕТКА, ОТСТУПЫ, РАМКИ И ЗАЩИТА ОТ РАЗРЫВОВ)
# =====================================================================

def set_cell_margins(cell, top_mm, bottom_mm, left_mm, right_mm):
    """Устанавливает внутренние отступы ячейки в миллиметрах (padding)."""
    def mm_to_dxa(mm): return int((mm / 25.4) * 1440)
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, val in zip(['top', 'bottom', 'left', 'right'], [top_mm, bottom_mm, left_mm, right_mm]):
        node = OxmlElement(f'w:{margin}')
        node.set(qn('w:w'), str(mm_to_dxa(val)))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, outer_style, inner_style, border_size_pt=1):
    """Настройка внешних и внутренних границ таблицы через OXML."""
    sz_val = str(int(border_size_pt * 8))
    tblPr = table._element.tblPr
    tblBorders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), outer_style)
        border.set(qn('w:sz'), sz_val)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)

    for border_name in ['insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), inner_style)
        border.set(qn('w:sz'), sz_val)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)

    tblPr.append(tblBorders)

# =====================================================================
# БЛОК 2: АНАЛИЗАТОР И ПЕРЕНОСЧИК СТИЛЕЙ XML (ДЛЯ КЛОНИРОВАНИЯ ФОНОВ/ЦВЕТОВ)
# =====================================================================

def extract_run_style_xml(run):
    """
    Вытаскивает все стили, включая цвет текста, цвет фона (Shading)
    и цвет маркера (Highlight) напрямую из разметки XML.
    """
    rPr = run._element.rPr
    style = {
        'bold': bool(run.bold),
        'italic': bool(run.italic),
        'underline': bool(run.underline),
        'strike': bool(run.font.strike),
        'color': None,
        'shd': None,
        'highlight': None,
        'font_name': None,
        'font_size': None
    }
    if rPr is not None:
        # Извлекаем точный цвет шрифта
        color_el = rPr.find(qn('w:color'))
        if color_el is not None:
            style['color'] = color_el.get(qn('w:val'))
        
        # Извлекаем фоновую заливку (Pink/Yellow Shading из вашего исходника)
        shd = rPr.find(qn('w:shd'))
        if shd is not None:
            style['shd'] = shd.get(qn('w:fill'))
            
        # Извлекаем маркер (Highlight)
        hl = rPr.find(qn('w:highlight'))
        if hl is not None:
            style['highlight'] = hl.get(qn('w:val'))
            
        # Название шрифта
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is not None:
            style['font_name'] = rFonts.get(qn('w:ascii')) or rFonts.get(qn('w:hAnsi'))
            
        # Размер шрифта
        sz = rPr.find(qn('w:sz'))
        if sz is not None:
            try:
                style['font_size'] = float(sz.get(qn('w:val'))) / 2.0
            except:
                pass
    return style

def apply_xml_style_to_run(run, style, global_config):
    """Применяет накопленные XML стили к новой ячейке Word-карточки."""
    rPr = run._element.get_or_add_rPr()
    
    if style['bold']: run.bold = True
    if style['italic']: run.italic = True
    if style['underline']: run.underline = True
    if style['strike']: run.font.strike = True
    
    # Цвет шрифта
    if style['color']:
        color_el = OxmlElement('w:color')
        color_el.set(qn('w:val'), style['color'])
        rPr.append(color_el)
        
    # Фоновая заливка (Shading)
    if style['shd']:
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), style['shd'])
        rPr.append(shd)
        
    # Выделение маркером (Highlight)
    if style['highlight']:
        hl = OxmlElement('w:highlight')
        hl.set(qn('w:val'), style['highlight'])
        rPr.append(hl)
        
    # Гарнитура и размер шрифта
    f_name = style['font_name'] or global_config['font_name']
    f_size = style['font_size'] or global_config['font_size']
    
    if f_name:
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), f_name)
        rFonts.set(qn('w:hAnsi'), f_name)
        rPr.append(rFonts)
        
    if f_size:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(f_size * 2)))
        rPr.append(sz)

# =====================================================================
# БЛОК 3: УМНЫЙ ОБРАБОТЧИК СЛОВАРНЫХ ПРАВИЛ И ВЫДЕЛЕНИЙ
# =====================================================================

def hex_to_rgb(hex_str):
    hex_str = hex_str.lstrip('#')
    if len(hex_str) != 6: return None
    return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))

def apply_rules_to_word(word, style, rules):
    """
    Применяет пользовательские правила (авто-выделение слов,
    кавычки, капс) поверх оригинального стиля.
    """
    new_style = style.copy()
    clean_word = word.strip(".,;:!?\"'«»()[]")
    
    # 1. Авто-выделение пользовательских слов
    if rules.get('custom_bold_words'):
        bold_list = [w.strip().lower() for w in rules['custom_bold_words'].split(',') if w.strip()]
        if clean_word.lower() in bold_list:
            new_style['bold'] = True
            if rules.get('custom_words_color_on'):
                new_style['color'] = rules['custom_words_color'].lstrip('#')
                
    # 2. Авто-выделение слов внутри кавычек «...» или "..."
    if rules.get('bold_quotes'):
        if any(q in word for q in ['«', '»', '"', '“', '”']):
            new_style['bold'] = True

    # 3. Базовые умные правила
    if rules.get('enable_smart_rules', False):
        word_clean = word.strip()
        if word_clean.isupper() and any(c.isalpha() for c in word_clean) and len(word_clean) > 1:
            if rules.get('caps_bold'): new_style['bold'] = True
            if rules.get('caps_italic'): new_style['italic'] = True
            if rules.get('caps_color_on'): new_style['color'] = rules['caps_color'].lstrip('#')
            
        if re.match(r'^[\dа-яА-Яa-zA-Z]+\)$', word_clean):
            if rules.get('list_bold'): new_style['bold'] = True
            if rules.get('list_italic'): new_style['italic'] = True
            if rules.get('list_color_on'): new_style['color'] = rules['list_color'].lstrip('#')
            
    return new_style

# =====================================================================
# БЛОК 4: ПАРСЕР ДОКУМЕНТА (С КЛОНИРОВАНИЕМ ФОРМАТОВ И СЖАТИЕМ РЯДОВ)
# =====================================================================

def parse_docx_file(file_bytes, rules, global_config):
    """
    Считывает структуру Word, бережно сохраняя родное XML форматирование,
    и оптимизирует пустые ряды, если включена компактность.
    """
    stream = []
    doc = Document(file_bytes)
    collapse_empty = global_config.get('collapse_empty_lines', False)
    
    for p in doc.paragraphs:
        text_stripped = p.text.strip()
        if not text_stripped:
            if not collapse_empty:
                stream.append(("\n", {'bold': False, 'italic': False, 'underline': False, 'strike': False, 'color': None, 'shd': None, 'highlight': None, 'font_name': None, 'font_size': None}))
            continue
            
        for run in p.runs:
            text = run.text
            if not text: continue
            
            orig_fmt = extract_run_style_xml(run)
            words = text.split(' ')
            for i, word in enumerate(words):
                if not word and i != len(words)-1:
                    stream.append((" ", orig_fmt))
                    continue
                if not word: continue
                
                final_fmt = apply_rules_to_word(word, orig_fmt, rules)
                suffix = " " if i < len(words) - 1 else ""
                stream.append((word + suffix, final_fmt))
                
        stream.append(("\n", {'bold': False, 'italic': False, 'underline': False, 'strike': False, 'color': None, 'shd': None, 'highlight': None, 'font_name': None, 'font_size': None}))
        
    # Сжатие рядов (удаление множественных переносов)
    if collapse_empty:
        filtered_stream = []
        prev_was_newline = False
        for text_chunk, fmt in stream:
            if text_chunk == "\n":
                if prev_was_newline:
                    continue
                prev_was_newline = True
            else:
                prev_was_newline = False
            filtered_stream.append((text_chunk, fmt))
        stream = filtered_stream
        
    return stream

# =====================================================================
# БЛОК 5: ЯДРО СБОРКИ DOCX (С КЛОНИРОВАНИЕМ И ЗАЩИТОЙ cantSplit)
# =====================================================================

def build_professional_docx(word_stream, config):
    """Собирает сетку карточек, перенося 1-в-1 все стили оформления."""
    cells_data = []
    current_cell =[]
    current_count = 0
    
    # Буфер для накопления текущего абзаца/пункта
    buffer =[]
    buffer_len = 0
    
    for text_chunk, fmt in word_stream:
        buffer.append((text_chunk, fmt))
        weight = 0 if text_chunk == "\n" else len(text_chunk)
        buffer_len += weight
        
        # Как только встречаем конец абзаца (перенос строки)
        if text_chunk == "\n":
            # Проверяем: если добавление этого абзаца превысит лимит, 
            # и в карточке УЖЕ есть текст -> закрываем карточку и начинаем новую
            if current_count + buffer_len > config['max_chars'] and current_count > 0:
                cells_data.append(current_cell)
                current_cell =[]
                current_count = 0
                
            # Кладем целиком накопленный абзац в ячейку
            current_cell.extend(buffer)
            current_count += buffer_len
            
            # Очищаем буфер для следующего пункта
            buffer =[]
            buffer_len = 0
            
    # Если в конце файла остался кусок текста без переноса строки
    if buffer:
        if current_count + buffer_len > config['max_chars'] and current_count > 0:
            cells_data.append(current_cell)
            current_cell = buffer
        else:
            current_cell.extend(buffer)
            
    if current_cell:
        cells_data.append(current_cell)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    
    section.left_margin = Mm(config['page_margin'])
    section.right_margin = Mm(config['page_margin'])
    section.top_margin = Mm(config['page_margin'])
    section.bottom_margin = Mm(config['page_margin'])

    num_rows = len(cells_data) // config['cols_num']
    table = doc.add_table(rows=num_rows, cols=config['cols_num'])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    set_table_borders(
        table, 
        outer_style=config['border_outer'], 
        inner_style=config['border_inner'], 
        border_size_pt=config['border_size']
    )
    
    for row_idx, row in enumerate(table.rows):
        row.height = Mm(config['row_height'])
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        
        # ЗАЩИТА: запрещаем строкам/карточкам разрываться между страницами в Word
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:cantSplit'))
        
        for col_idx, cell in enumerate(row.cells):
            cell.width = Mm(config['col_width'])
            set_cell_margins(cell, config['pad_mm'], config['pad_mm'], config['pad_mm'], config['pad_mm'])
            
            data_idx = row_idx * config['cols_num'] + col_idx
            if data_idx >= len(cells_data): continue
            
            data = cells_data[data_idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if config['justify'] else WD_ALIGN_PARAGRAPH.LEFT
            
            p.paragraph_format.line_spacing = config['line_spacing']
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            
            if not data: continue
            
            for text_chunk, fmt_dict in data:
                if text_chunk == "\n":
                    p.add_run().add_break()
                else:
                    run = p.add_run(text_chunk)
                    apply_xml_style_to_run(run, fmt_dict, config)
    
    target = BytesIO()
    doc.save(target)
    target.seek(0)
    return target

# =====================================================================
# БЛОК 6: РУСИФИЦИРОВАННЫЙ ИНТЕРФЕЙС STREAMLIT С ПОДПИСЯМИ
# =====================================================================

st.set_page_config(page_title="PRO Генератор Шпор", page_icon="✂️", layout="wide")

st.markdown("""
<style>
    .main-header { font-size: 2.3rem; font-weight: 800; color: #1E1E1E; margin-bottom: 0px;}
    .sub-header { font-size: 1.05rem; color: #555555; margin-bottom: 25px;}
    .st-expander { border-radius: 8px !important; border: 1px solid #ddd !important; }
</style>
""", unsafe_allow_html=True)

st.markdown('<div class="main-header">✂️ PRO Генератор Шпор (Клонирование 1-в-1)</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-header">Клонирует исходное форматирование Word-файла (цветной текст, розовые/желтые маркеры фона) и собирает ровную сетку карточек.</div>', unsafe_allow_html=True)

col_settings, col_main = st.columns([1.3, 2.2], gap="large")

with col_settings:
    st.write("### 🎛️ Параметры Сетки")
    
    with st.expander("📄 Разметка и Размеры карточек", expanded=True):
        cols_num = st.number_input("Количество колонок", min_value=1, max_value=10, value=4)
        page_margin = st.number_input("Поля страницы А4 (мм)", min_value=0.0, max_value=50.0, value=10.0, step=1.0)
        col_width = st.number_input("Ширина карточки (мм)", min_value=30.0, max_value=100.0, value=51.0, step=0.5)
        row_height = st.number_input("Высота карточки (мм)", min_value=50.0, max_value=150.0, value=88.3, step=0.1)
        pad_mm = st.number_input("Внутренний отступ текста от рамки (мм)", min_value=0.0, max_value=10.0, value=1.5, step=0.1)

    with st.expander("⚡ Оптимизация и Компактность", expanded=True):
        collapse_empty_lines = st.checkbox("🗜️ Сомкнуть ряды (Сжатие пустот)", value=True, help="Убирает повторяющиеся пустые строки и пустые абзацы из исходного файла, делая шпоры максимально плотными.")

    with st.expander("🔤 Параметры Шрифтов по умолчанию", expanded=False):
        font_preset = st.selectbox("Семейство шрифта", ["Raleway", "Times New Roman", "Arial", "Calibri", "Свой шрифт..."])
        font_custom = ""
        if font_preset == "Свой шрифт...":
            font_custom = st.text_input("Название шрифта:", "Montserrat")
        final_font = font_custom if font_preset == "Свой шрифт..." else font_preset
        
        font_size = st.number_input("Базовый кегль (pt)", min_value=3.0, max_value=14.0, value=5.0, step=0.5)
        justify = st.checkbox("Выравнивание по ширине (Justify)", value=False)
        line_spacing = st.number_input("Базовый межстрочный интервал", min_value=0.5, max_value=2.0, value=0.92, step=0.01)
        max_chars = st.slider("Лимит символов в одной карточке", 1000, 4000, 2100, step=50)

    with st.expander("🔲 Границы и Рамки", expanded=False):
        b1, b2 = st.columns(2)
        border_options = ["single", "dashed", "dotted", "double", "nil"]
        with b1: border_outer = st.selectbox("Внешняя рамка", border_options, index=0)
        with b2: border_inner = st.selectbox("Внутренняя сетка (Линии реза)", border_options, index=1)
        border_size = st.number_input("Толщина линий рамок (pt)", min_value=0.5, max_value=5.0, value=1.0, step=0.5)

    st.write("### 🎯 Доп. Авто-выделения")
    
    with st.expander("🎨 Выделение конкретных слов и кавычек", expanded=False):
        bold_quotes = st.checkbox("Выделять слова в кавычках («...», \"...\") полужирным", value=True)
        custom_bold_words = st.text_area("Список слов для авто-выделения полужирным (через запятую):", help="Например: Обоснуйте, человек, закон, истина")
        custom_words_color_on = st.checkbox("Красить эти слова в свой цвет", value=False)
        custom_words_color = "#FF0000"
        if custom_words_color_on:
            custom_words_color = st.color_picker("Цвет для авто-выделяемых слов", "#FF0000")

    # Спящие умные правила (выключены по умолчанию)
    enable_smart_rules = st.checkbox("🤖 Активировать умную перезапись стилей", value=False)
    with st.expander("Настройки умных стилей", expanded=enable_smart_rules):
        st.markdown("**Для КАПСЛОКА:**")
        r1, r2, r3 = st.columns(3)
        caps_bold = r1.checkbox("Жирный", value=True, disabled=not enable_smart_rules)
        caps_italic = r2.checkbox("Курсив", value=False, disabled=not enable_smart_rules)
        caps_color_on = r3.checkbox("Цвет", value=False, disabled=not enable_smart_rules)
        caps_color = st.color_picker("Цвет КАПСА", "#FF0000", disabled=not enable_smart_rules)
        
        st.markdown("**Для нумерации списков (1), а):**")
        l1, l2, l3 = st.columns(3)
        list_bold = l1.checkbox("Жирный ", value=True, key="lb", disabled=not enable_smart_rules)
        list_italic = l2.checkbox("Курсив ", value=False, key="li", disabled=not enable_smart_rules)
        list_color_on = l3.checkbox("Цвет ", value=False, key="lc", disabled=not enable_smart_rules)
        list_color = st.color_picker("Цвет списков", "#0000FF", disabled=not enable_smart_rules)

    config = {
        'cols_num': cols_num, 'page_margin': page_margin, 'col_width': col_width, 
        'row_height': row_height, 'pad_mm': pad_mm, 'border_outer': border_outer, 
        'border_inner': border_inner, 'border_size': border_size,
        'font_name': final_font, 'font_size': font_size, 'justify': justify, 
        'line_spacing': line_spacing, 'max_chars': max_chars,
        'collapse_empty_lines': collapse_empty_lines,
        'rules': {
            'enable_smart_rules': enable_smart_rules,
            'bold_quotes': bold_quotes,
            'custom_bold_words': custom_bold_words,
            'custom_words_color_on': custom_words_color_on,
            'custom_words_color': custom_words_color,
            'caps_bold': caps_bold, 'caps_italic': caps_italic, 
            'caps_color_on': caps_color_on, 'caps_color': caps_color,
            'list_bold': list_bold, 'list_italic': list_italic, 
            'list_color_on': list_color_on, 'list_color': list_color
        }
    }

with col_main:
    st.write("### 📁 Клонирование из Вашего Word-файла")
    st.info("Программа выполнит парсинг вашего файла, перенесёт все выделенные цветным маркером абзацы и шрифты в новую упорядоченную сетку карточек-шпор.")
    
    uploaded_file = st.file_uploader("Загрузите файл .docx", type=["docx"])
    
    if st.button("🚀 Сгенерировать шпоры по сетке", use_container_width=True, type="primary"):
        if uploaded_file is None:
            st.error("❌ Загрузите документ для генерации.")
        else:
            with st.spinner('Анализ документа, перенос заливок, шрифтов и сборка карточек...'):
                try:
                    word_stream = parse_docx_file(uploaded_file, config['rules'], config)
                    result_doc = build_professional_docx(word_stream, config)
                    
                    st.success("✅ Карточки сгенерированы идеально с сохранением всех цветов фона и стилей!")
                    st.download_button(
                        label="📥 Скачать идеальные шпоры (.docx)",
                        data=result_doc,
                        file_name=f"Cards_1to1_Cloned_{uploaded_file.name}",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                except Exception as e:
                    st.error(f"❌ Произошла ошибка: {e}")

st.markdown("---")
st.caption("Клонирование Word 1-в-1. cantSplit на уровне ячеек гарантирует отсутствие разрывов.")
